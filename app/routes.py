from flask import Blueprint, render_template, request, jsonify, current_app, Response
from io import BytesIO
from gtts import gTTS
import os
from werkzeug.utils import secure_filename
from docx import Document
import PyPDF2
import traceback
import docx2txt  # For text extraction

main = Blueprint('main', __name__)

ALLOWED_EXTENSIONS = {'txt', 'docx', 'pdf'}

def allowed_file(filename):
    """Check if the file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def append_to_docx(input_file, filename):
    """Append uploaded file content to scraping_results.docx safely."""
    # If scraping_results.docx exists, extract its content with docx2txt to avoid depth issues
    existing_content = ""
    if os.path.exists("scraping_results.docx"):
        existing_content = docx2txt.process("scraping_results.docx")

    # Create a new document to avoid inheriting deep XML structure
    target_doc = Document()

    # Add existing content if any
    if existing_content.strip():
        target_doc.add_paragraph(existing_content)
        target_doc.add_paragraph().add_run().add_break()

    # Add new content headers
    target_doc.add_paragraph(f"Page {len(target_doc.paragraphs) + 1}")
    target_doc.add_paragraph(f"URL: Uploaded File - {filename}")
    target_doc.add_paragraph("Meta Description: Uploaded document content")
    target_doc.add_paragraph("Main Content")

    # Append new content based on file type
    if filename.endswith('.txt'):
        content = input_file.read().decode('utf-8')
        target_doc.add_paragraph(content)
    
    elif filename.endswith('.docx'):
        content = docx2txt.process(input_file)
        target_doc.add_paragraph(content)
    
    elif filename.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(input_file)
        text_content = []
        for page in pdf_reader.pages:
            text_content.append(page.extract_text())
        target_doc.add_paragraph('\n'.join(text_content))

    target_doc.add_paragraph("Contact Information")
    target_doc.save("scraping_results.docx")

@main.route('/')
def index():
    """Render the main index page."""
    return render_template('index.html')

@main.route('/listen')
def listen_page():
    """Render the listen page."""
    return render_template('listen.html')

@main.route('/listen', methods=['POST'])
def listen():
    """Handle audio response generation."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Invalid request format'}), 400
            
        user_message = data.get('message', '')
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400

        if not hasattr(current_app, 'chatbot'):
            return jsonify({'error': 'System error'}), 500

        chatbot = current_app.chatbot
        response_text = ""
        similar_docs = []

        # Handle conversation state
        if chatbot.first_interaction:
            response_text = "Hello! I'm SAM, your student consultant. What's your name?"
            chatbot.first_interaction = False
        elif chatbot.user_name is None:
            chatbot.user_name = user_message
            response_text = f"Nice to meet you, {user_message}! How can I help you today?"
        else:
            response_text, similar_docs = chatbot.get_response(user_message)

        # Generate audio
        tts = gTTS(text=response_text, lang='en')
        audio_buffer = BytesIO()
        tts.write_to_fp(audio_buffer)
        audio_buffer.seek(0)

        return Response(
            audio_buffer,
            mimetype="audio/mpeg",
            headers={
                "Content-Disposition": "inline; filename=response.mp3",
                "X-Response-Text": response_text
            }
        )

    except Exception as e:
        current_app.logger.error(f"Error in /listen: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@main.route('/documents')
def documents():
    """Render the documents page."""
    return render_template('documents.html')

@main.route('/upload-document', methods=['POST'])
def upload_document():
    """Handle document upload and processing."""
    try:
        if 'document' not in request.files:
            return jsonify({'error': 'No document part'}), 400
        
        file = request.files['document']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed'}), 400
        
        filename = secure_filename(file.filename)
        append_to_docx(file, filename)
        
        if os.path.exists('embeddings_cache.pkl'):
            os.remove('embeddings_cache.pkl')
        if os.path.exists('document_hash.txt'):
            os.remove('document_hash.txt')
        
        current_app.chatbot.doc_store.load_docx("scraping_results.docx")
        current_app.chatbot.doc_store.create_embeddings()
        
        return jsonify({'message': 'Document uploaded successfully'})
    except Exception as e:
        current_app.logger.error(f"Upload error: {str(e)}")
        return jsonify({'error': f"Failed to process document: {str(e)}"}), 500

@main.route('/chat', methods=['POST'])
def chat():
    """Handle chat interactions."""
    try:
        data = request.get_json()
        user_message = data.get('message', '')
        
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400
        
        chatbot = current_app.chatbot
        response_text = ""
        similar_docs = []

        # Handle conversation state
        if chatbot.first_interaction:
            response_text = "Hello! I'm SAM, your student consultant. What's your name?"
            chatbot.first_interaction = False
        elif chatbot.user_name is None:
            chatbot.user_name = user_message
            response_text = f"Nice to meet you, {user_message}! How can I help you today?"
        else:
            response_text, similar_docs = chatbot.get_response(user_message)

        return jsonify({
            'response': response_text,
            'similar_docs': similar_docs
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500